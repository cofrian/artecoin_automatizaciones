#!/usr/bin/env python
# -*- coding: utf-8 -*-

import argparse, json, re, tempfile
from pathlib import Path
from typing import Dict, List, Tuple, Optional

from docxtpl import DocxTemplate
from jinja2 import Environment, DebugUndefined
from docx.shared import Cm, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Pillow (recomendado) para leer tamaño real de imagen y reparar formatos
try:
    from PIL import Image
    PIL_OK = True
except Exception:
    PIL_OK = False

# ================== PARÁMETROS DE MAQUETACIÓN ==================
# Medidas de la CELDA donde irán las fotos
TARGET_CELL_W_CM = 18.0   # ancho total de la celda
TARGET_CELL_H_CM = 22.0   # alto total de la celda

# Margen perimetral dentro de esa celda y separación entre celdas de foto
MARGIN_CM       = 0.6     # margen alrededor (arriba/abajo/izquierda/derecha)
GUTTER_CM       = 0.4     # separación entre imágenes (horizontal/vertical)
CAPTION_H_CM    = 0.8     # alto reservado para el pie de foto

# Peso extra para la fila “grande” en layouts 3 (2+1) y 5 (2,2,1)
ROW_WEIGHT_SINGLE_BIG = 1.4

# ===============================================================

# --------- util plantilla ---------
def _load_json(p: Path) -> Dict:
    return json.loads(Path(p).read_text(encoding="utf-8"))

def _delete_paragraph(paragraph):
    p = paragraph._p
    p.getparent().remove(p)
    paragraph._element = None

def _find_paragraphs_with_marker(doc, marker: str):
    hits = []
    for p in doc.paragraphs:
        if marker in p.text:
            hits.append(p)
    # También dentro de tablas:
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if marker in p.text:
                        hits.append(p)
    return hits

def _set_tbl_width_pct(tbl, pct_0_100=100):
    """Establece ancho de tabla en % de su contenedor (robusto)."""
    # Evita errores de tblW inexistente
    tbl_el = tbl._tbl
    tblPr = tbl_el.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl_el.insert(0, tblPr)
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:type'), 'pct')
    # Word usa 5000 = 100%
    tblW.set(qn('w:w'), str(int(pct_0_100 * 50)))
    # Centrar
    try:
        tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        pass

def _safe_image_path(path_str: str) -> str:
    """
    Devuelve una ruta utilizable. Si la imagen rompe (EXIF/HEIC/WEBP),
    re-codifica a PNG temporal (si Pillow está disponible).
    """
    p = Path(path_str)
    try:
        with open(p, "rb") as fh:
            fh.read(32)
        return str(p)
    except Exception:
        pass

    if not PIL_OK:
        return str(p)

    try:
        img = Image.open(path_str).convert("RGB")
        tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
        img.save(tmp.name, format="PNG")
        return tmp.name
    except Exception:
        return str(p)

def _image_size_cm(path_str: str) -> Optional[Tuple[float, float]]:
    """Devuelve (ancho_cm, alto_cm) de la imagen si Pillow está disponible."""
    if not PIL_OK:
        return None
    try:
        with Image.open(path_str) as im:
            w_px, h_px = im.size
            # Asumimos 96 dpi si no hay info; si hay info DPI, la usamos
            dpi = im.info.get("dpi", (96, 96))
            xdpi = dpi[0] if isinstance(dpi, (list, tuple)) and dpi else 96
            ydpi = dpi[1] if isinstance(dpi, (list, tuple)) and len(dpi) > 1 else 96
            # px -> pulgadas -> cm
            w_cm = (w_px / (xdpi or 96)) * 2.54
            h_cm = (h_px / (ydpi or 96)) * 2.54
            return (w_cm, h_cm)
    except Exception:
        return None

def _insert_caption(cell, text: str, style_name: str = "PieFotoCorp"):
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    # estilo si existe; fallback si no
    try:
        p.style = style_name
    except Exception:
        run.font.size = Pt(8)

def _fill_foto_cell(cell, item: Dict, width_cm: float, height_cm: float):
    """
    Limpia la celda y mete imagen centrada + pie.
    Ajusta por ancho y alto disponibles manteniendo proporción.
    """
    for par in list(cell.paragraphs):
        _delete_paragraph(par)

    img_path_raw = item["path"]
    img_path = _safe_image_path(img_path_raw)
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()

    # Calcula tamaño destino respetando caja (width_cm x height_cm - CAPTION_H_CM)
    pic_w = width_cm
    pic_h = height_cm

    # Si conocemos tamaño real, ajustamos por ancho/alto
    real = _image_size_cm(img_path)
    if real:
        rw, rh = real
        if rw > 0 and rh > 0:
            scale = min(pic_w / rw, pic_h / rh)
            pic_w = max(0.5, rw * scale)
            pic_h = max(0.5, rh * scale)
    # Insertamos usando ancho; Word respetará proporción
    try:
        r.add_picture(img_path, width=Cm(pic_w))
    except Exception:
        # Reparo final a PNG si aún falla
        if PIL_OK:
            try:
                from PIL import Image
                im = Image.open(img_path).convert("RGB")
                tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
                im.save(tmp.name, format="PNG")
                r.add_picture(tmp.name, width=Cm(pic_w))
            except Exception:
                p.add_run("⚠ No se pudo insertar la imagen")
        else:
            p.add_run("⚠ No se pudo insertar la imagen")

    caption = item.get("id") or item.get("name") or ""
    if caption:
        _insert_caption(cell, f"FOTO: {caption}")

# ------------------ LAYOUTS ------------------
def _compute_layout(fotos: List[Dict]) -> List[List[int]]:
    """
    Devuelve índices por filas según número de fotos.
    Ej.: 3 -> [[0,1],[2]]  ; 5 -> [[0,1],[2,3],[4]]
    """
    n = len(fotos)
    idx = list(range(n))
    if n <= 0:
        return []
    if n == 1:
        return [[0]]
    if n == 2:
        return [[0], [1]]
    if n == 3:
        return [[0, 1], [2]]          # 2 arriba, 1 abajo grande
    if n == 4:
        return [[0, 1], [2, 3]]       # 2x2
    if n == 5:
        return [[0, 1], [2, 3], [4]]  # 2,2,1 (última grande)
    if n == 6:
        return [[0, 1, 2], [3, 4, 5]] # 3x2
    # >=7 → filas de 3
    out = []
    i = 0
    while i < n:
        out.append(idx[i:i+3])
        i += 3
    return out

def _row_weights(layout: List[List[int]]) -> List[float]:
    """Mayor peso a la fila única de los layouts 3 (2+1) y 5 (2,2,1)."""
    n_rows = len(layout)
    if n_rows == 0:
        return []
    if n_rows == 2 and len(layout[1]) == 1:
        return [1.0, ROW_WEIGHT_SINGLE_BIG]  # 3 fotos
    if n_rows == 3 and len(layout[2]) == 1:
        return [1.0, 1.0, ROW_WEIGHT_SINGLE_BIG]  # 5 fotos
    return [1.0] * n_rows

def _build_table_in_cell(cell, fotos: List[Dict]):
    if not fotos:
        return

    # Área útil dentro de la celda (descontar márgenes)
    inner_w = max(1e-3, TARGET_CELL_W_CM - 2 * MARGIN_CM)
    inner_h = max(1e-3, TARGET_CELL_H_CM - 2 * MARGIN_CM)

    layout = _compute_layout(fotos)
    n_rows = len(layout)
    weights = _row_weights(layout)
    sum_w = sum(weights) if weights else 1.0

    # Altura disponible para imágenes + pies (y gutters verticales)
    total_gutters_v = max(0, n_rows - 1) * GUTTER_CM
    avail_h_total = max(0.1, inner_h - total_gutters_v)

    # Crea tabla interna
    max_cols = max(len(row) for row in layout)
    tbl = cell.add_table(rows=n_rows, cols=max_cols)
    _set_tbl_width_pct(tbl, 100)

    # Opcional: aplica un estilo si existe
    try:
        if "TablaFotosCorp" in [s.name for s in cell._tc.part.document.styles]:
            tbl.style = "TablaFotosCorp"
    except Exception:
        pass

    for ri, row_idx in enumerate(layout):
        row_cols = len(row_idx)
        # Alto de fila ponderado
        row_h = (weights[ri] / sum_w) * avail_h_total
        # Caja vertical para la imagen en esa fila (restar el pie)
        pic_h = max(0.1, row_h - CAPTION_H_CM)

        # Ancho disponible en esa fila (descontando gutters horizontales)
        total_gutters_h = max(0, row_cols - 1) * GUTTER_CM
        row_inner_w = max(0.1, inner_w - total_gutters_h)
        # Ancho de celda para cada imagen
        pic_w_each = row_inner_w / max(1, row_cols)

        # Si la fila tiene menos columnas que max_cols, fusionamos celdas sobrantes
        if row_cols < max_cols:
            # Fusionar desde la col 0 hasta col (max_cols-1) si row_cols=1 (imagen grande)
            if row_cols == 1:
                cell0 = tbl.cell(ri, 0)
                last = tbl.cell(ri, max_cols - 1)
                cell0.merge(last)
                # Rellenar en la celda fusionada
                foto = fotos[row_idx[0]]
                _fill_foto_cell(cell0, foto, width_cm=pic_w_each * max_cols, height_cm=pic_h)
                # Limpia celdas que quedaron (por seguridad)
                for ci in range(1, max_cols):
                    for p in list(tbl.cell(ri, ci).paragraphs):
                        _delete_paragraph(p)
                continue

        # Relleno normal cuando no hay fusión o hay varias celdas
        for ci in range(max_cols):
            c = tbl.cell(ri, ci)
            if ci < row_cols:
                foto = fotos[row_idx[ci]]
                _fill_foto_cell(c, foto, width_cm=pic_w_each, height_cm=pic_h)
            else:
                # Celda vacía -> limpiar
                for p in list(c.paragraphs):
                    _delete_paragraph(p)
                c.add_paragraph("")

# ------------------ BÚSQUEDA / REEMPLAZO DE MARCADORES ------------------
def _norm_fotos_struct(lst):
    """
    Admite 'fotos' (lista de dicts con 'path','id','name') o 'fotos_por_filas' (legacy).
    Retorna lista plana de dicts con 'path','id','name'.
    """
    if isinstance(lst, list) and lst and isinstance(lst[0], dict) and "path" in lst[0]:
        return lst
    flat = []
    for fila in lst or []:
        for f in fila:
            flat.append({"path": f.get("path",""), "name": f.get("name",""), "id": f.get("name","")})
    return flat

def _gather_foto_lookup(ctx: Dict) -> Dict[str, List[Dict]]:
    """
    Índice clave->lista_fotos. Claves reconocidas:
      FOTOS_CENTRO_{centro.id}
      FOTOS_EDIFICIO_{e.id}
      FOTOS_DEPENDENCIA_{d.id}
      FOTOS_ENVOL_{it.id}, FOTOS_SISTCC_{it.id}, FOTOS_CLIMA_{it.id},
      FOTOS_EQHORIZ_{it.id}, FOTOS_ELEVA_{it.id}, FOTOS_ILUM_{it.id}, FOTOS_OTROSEQ_{it.id}
    """
    out = {}
    cent = ctx.get("centro", {})
    f_c = _norm_fotos_struct(cent.get("fotos") or cent.get("fotos_por_filas"))
    if f_c:
        out[f"FOTOS_CENTRO_{cent.get('id','')}"] = f_c

    for e in ctx.get("edif", []):
        f_e = _norm_fotos_struct(e.get("fotos") or e.get("fotos_por_filas"))
        if f_e:
            out[f"FOTOS_EDIFICIO_{e.get('id','')}"] = f_e

        for d in e.get("dependencias", []):
            f_d = _norm_fotos_struct(d.get("fotos") or d.get("fotos_por_filas"))
            if f_d:
                out[f"FOTOS_DEPENDENCIA_{d.get('id','')}"] = f_d

        for key, tag in [
            ("envolventes", "ENVOL"),
            ("sistemas_cc", "SISTCC"),
            ("equipos_clima", "CLIMA"),
            ("equipos_horiz", "EQHORIZ"),
            ("elevadores", "ELEVA"),
            ("iluminacion", "ILUM"),
            ("otros_equipos", "OTROSEQ"),
        ]:
            for it in e.get(key, []):
                f_it = _norm_fotos_struct(it.get("fotos") or it.get("fotos_por_filas"))
                if f_it:
                    out[f"FOTOS_{tag}_{it.get('id','')}"] = f_it
    return out

def _replace_foto_markers(doc, ctx: Dict):
    lookup = _gather_foto_lookup(ctx)
    for key, fotos in lookup.items():
        marker = f"[[{key}]]"
        hits = _find_paragraphs_with_marker(doc, marker)
        for p in hits:
            # Borrar el texto del marcador pero quedarnos en esa celda
            p.text = p.text.replace(marker, "")
            # El parent de un párrafo en celda es el objeto Cell
            cell = getattr(p, "_parent", None)
            # Si no estamos en celda, creamos tabla 1x1 justo "después"
            if cell is None or not hasattr(cell, "add_table"):
                tbl = doc.add_table(rows=1, cols=1)
                _set_tbl_width_pct(tbl, 100)
                _build_table_in_cell(tbl.cell(0, 0), fotos)
            else:
                _build_table_in_cell(cell, fotos)

# ----------------- RENDER -----------------
def render_centro(json_file: Path, plantilla: Path, out_dir: Path) -> Path:
    ctx = _load_json(json_file)

    env = Environment(undefined=DebugUndefined)
    doc = DocxTemplate(str(plantilla))
    # Primera pasada: sólo texto Jinja (no fotos)
    doc.render(ctx, jinja_env=env)

    # Segunda pasada: sustituir marcadores [[FOTOS_*_ID]]
    _replace_foto_markers(doc, ctx)

    out_dir.mkdir(parents=True, exist_ok=True)
    out = out_dir / (Path(json_file).stem + ".docx")
    doc.save(str(out))
    return out

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--plantilla", required=True)
    ap.add_argument("--json-dir", required=True)
    ap.add_argument("--out-dir", required=True)
    args = ap.parse_args()

    tpl = Path(args.plantilla)
    jdir = Path(args.json_dir)
    out = Path(args.out_dir)

    jsons = sorted(jdir.glob("*.json"))
    print(f"Renderizando {len(jsons)} centro(s)…")
    for jf in jsons:
        of = render_centro(jf, tpl, out)
        print(f"  ✓ {of.name}")
    print(f"[OK] Documentos guardados en: {out.resolve()}")

if __name__ == "__main__":
    main()
